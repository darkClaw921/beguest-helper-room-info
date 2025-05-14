from pprint import pprint
from urllib.parse import quote,unquote

import requests
from test import encode_yandex_url
from dotenv import load_dotenv
import os

load_dotenv()

IP_CHAT_ROOM=os.getenv('IP_CHAT_ROOM')
WEBHOOK_API_TOKEN=os.getenv('WEBHOOK_API_TOKEN')

def get_path_to_file_instruction(url, returnIsUrl=False)->list[str]:
    """
    Получает путь к файлу из URL
    in: https://docs.yandex.ru/docs/view?url=ya-disk-public%3A%2F%mFJ6bpmRеса%2C%20Инструкции%2FМалпер.%202%20кор.7-237).docx&name=Малевич(Трамвайный%20пер.%202%20кор.7-237).docx 
    out:Адреса, Инструкции/8 марта 204Д - 116 (16) .docx
    returnIsUrl:True - возвращает URL без кодировки
    https://docs.yandex.ru/docs/view?url=ya-disk-public://Np/fsAs9J/KnY3t95176lSoayq/J6bpmRyoXnDag==:/Адреса, Инструкции/Малевич(Трамвайный пер. 2 кор.7-237).docx&name=Малевич(Трамвайный пер. 2 кор.7-237).docx
    """
    urls=url.split(',')
    # Разбиваем URL на части по параметрам
    prepare_urls=[]
    for url in urls:
        if returnIsUrl:
            encoded_url=encode_yandex_url(url)
            # print(encoded_url)
            prepare_urls.append(encoded_url)
            continue
            
            
            
        # Логика для returnIsUrl=False (существующая логика)
        url_parts = url.split('?')[1].split('&')
        
        # Получаем значение параметра url
        file_url = ''
        for part in url_parts:
            if part.startswith('url='):
                file_url = part.replace('url=', '')
                break
        
        # Декодируем URL-encoded строку
        decoded_url = file_url.replace('%20', ' ').replace('%2F', '/').replace('%3A', ':').replace('%2C', ',')
        
        # Извлекаем путь после последнего двоеточия
        path = decoded_url.split(':')[-1]
        prepare_urls.append(path)
        
    # pprint(prepare_urls)
    return prepare_urls

def get_path_to_file_info(url, returnIsUrl=False)->list[str]:
    """
    Получает путь к файлу из URL
    in: https://disk.yandex.ru/d/GdB/ИНФО%20ПО%/0204д%20-%20116%20(16%20этаж)/Где%20мусорка.mov
    out:/ИНФО ПО КВАРТИРАМ/8 марта 204д - 116 (16 этаж)/Где мусорка.mov
    returnIsUrl:True - возвращает URL без кодировки
    'https://disk.yandex.ru/d/GFQ/ИНФО ПО КВАРТИРАМ/8 марта 204д - 116 (16 этаж)/Где мусорка.mov'
    """
    try:
        urls=url.split(',')
    except:
        urls=url

    prepare_urls=[]
    for url in urls:
        if returnIsUrl:
            decoded_url = url.replace('%20', ' ').replace('%2F', '/').replace('%2C', ',')
            # print(decoded_url)
            encoded_url=prepare_url_encoding(decoded_url)
            # print(encoded_url)
            prepare_urls.append(encoded_url)
            continue
            
        path = url.split('/d/')[1]
        path = '/' + '/'.join(path.split('/')[1:])
        decoded_path = path.replace('%20', ' ').replace('%2F', '/').replace('%2C', ',')
        # print(decoded_path)
        prepare_urls.append(decoded_path)
    
    # print(prepare_urls)
    return prepare_urls

def prepare_url_encoding(url):
    """
    Преобразует декодированный URL в корректно закодированный URL
    in: https://disk.yandex.ru/d/SFQ/ИНФО ПО КВАРТИРАМ/Готвальда 24 корп.4 -103/Где мусорка.mov
    out: https://disk.yandex.ru/d/5N%98%D0%A0%D0%90%D0%9C/%D0%93%D0%03/%D0%93%D0%B4%D0%B5%20%D0%BC%D1%83%D1%81%D0%BE%D1%80%D0%BA%D0%B0.mov
    """
    # Сначала декодируем URL если он уже содержит кодированные части
    decoded_url = url.replace('%20', ' ').replace('%2F', '/').replace('%3A', ':').replace('%2C', ',').replace('%3D', '=')
    
    # Разделяем URL на части
    parts = decoded_url.split('/')
    
    # Кодируем только части пути после домена (первые 3 части оставляем без кодирования)
    encoded_parts = []
    for i, part in enumerate(parts):
        if i < 3:
            encoded_parts.append(part)
        else:
            encoded_parts.append(quote(part, safe=''))
    
    # Собираем URL обратно
    encoded_url = '/'.join(encoded_parts)
    
    return encoded_url




def extract_filenames(urls):
    filenames = []
    for url in urls:
        # Извлекаем часть после последнего '/'
        encoded_filename = url.split('/')[-1]
        # Декодируем URL-спецсимволы (например, %20 -> пробел)
        decoded_filename = unquote(encoded_filename)
        # Разделяем имя файла и расширение по последней точке
        if '.' in decoded_filename:
            name_parts = decoded_filename.rsplit('.', 1)
            filename_without_extension = name_parts[0]
        else:
            filename_without_extension = decoded_filename
        filenames.append(filename_without_extension)
    return filenames


mapping_deal_status={
    'C7:PREPARATION': 'Гость заехал',
    'C7:UC_3EBBY1': 'За 15 мин ничего не прислал',
    'C7:PREPAYMENT_INVOICE': 'Проверка оплаты из бота (если гость отправил скрин платежа)',

}


# mapping_deal_status={
#     '14': 'Заключен',
#     '13': 'Сделка отменена',
#     '12': 'Сделка отменена',
#     '11': 'Сделка отменена',
#     '10': 'Сделка отменена',
# }

def extract_text_from_docx(file_path):
    """
    Извлекает весь текст из docx файла
    
    Args:
        file_path (str): Путь к docx файлу
        
    Returns:
        str: Весь текст из документа
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Библиотека python-docx не установлена. Установите её командой 'uv add python-docx'")
    
    try:
        doc = Document(file_path)
        full_text = []
        
        # Извлекаем текст из каждого параграфа
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"Ошибка при извлечении текста: {str(e)}"
    

def convert_heic_to_jpg(input_path, output_path=None):
    """
    Конвертирует изображение из формата HEIC в формат JPG используя только Python библиотеки
    
    Args:
        input_path (str): Путь к HEIC файлу
        output_path (str, optional): Путь для сохранения JPG файла. 
                                    Если не указан, будет создан файл с тем же именем, но с расширением .jpg
    
    Returns:
        str: Путь к созданному JPG файлу или сообщение об ошибке
    """
    # Если путь для сохранения не указан, создаем его из исходного пути
    if output_path is None:
        output_path = input_path.rsplit('.', 1)[0] + '.jpg'
    
    try:
        # Пытаемся импортировать библиотеку pillow_heif
        try:
            import pillow_heif
            from PIL import Image
        except ImportError:
            return "Ошибка: Необходимые библиотеки не установлены. Установите их командой: pip install pillow-heif Pillow"
        
        # Регистрируем HEIF декодер в Pillow
        pillow_heif.register_heif_opener()
        
        # Открываем HEIC файл через Pillow (работает благодаря зарегистрированному декодеру)
        img = Image.open(input_path)
        
        # Если изображение в режиме RGBA, конвертируем в RGB
        if img.mode == 'RGBA':
            img = img.convert('RGB')
            
        # Сохраняем в формате JPEG
        img.save(output_path, format='JPEG', quality=95)
        
        return output_path
        
    except Exception as e:
        return f"Ошибка при конвертации HEIC в JPG: {str(e)}"
#  Примеры команд curl для работы с webhook API чата

# # 1. Отправка сообщения от менеджера клиенту
# # Этот запрос имитирует отправку сообщения от менеджера клиенту через внешнюю систему
# curl -X POST "http://localhost:8000/api/webhook/send-message" \
#   -H "Content-Type: application/json" \
#   -d '{
#     "telegram_id": 123456789, 
#     "text": "Здравствуйте! Это сообщение от менеджера через API.", 
#     "token": "your-webhook-api-token"
#   }'

# # 2. Отправка сообщения от клиента менеджеру
# # Этот запрос имитирует отправку сообщения от клиента в систему
# # curl -X POST "http://localhost:8000/api/webhook/client-message" \
# curl -X POST "http://localhost:8000/api/webhook/client-message" \
#   -H "Content-Type: application/json" \
#   -d '{
#     "telegram_id": 400923372, 
#     "text": "Здравствуйте! Это сообщение от клиента через API.", 
#     "token": "your-secret-api-token-here", 
#     "first_name": "Иван", 
#     "last_name": "Иванов", 
#     "username": "ivanov"
#   }'

# # 3. Отправка сообщения от клиента менеджеру без дополнительных данных
# # В этом случае будут использованы значения по умолчанию или данные из существующей записи пользователя
# curl -X POST "http://localhost:8000/api/webhook/client-message" \
#   -H "Content-Type: application/json" \
#   -d '{
#     "telegram_id": 123456789, 
#     "text": "Это еще одно сообщение от клиента.", 
#     "token": "your-webhook-api-token"
#   }'

# Примечание: замените "your-webhook-api-token" на актуальный токен из .env файла
# и при необходимости измените URL, если ваш сервер запущен на другом порту или адресе. 
def send_message_to_client(telegram_id, text):
    """Отправляет сообщение пользователю в телеграмм"""
    url=f'http://{IP_CHAT_ROOM}/api/webhook/send-message'
    data={
        'telegram_id': telegram_id,
        'text': text,
        'token': WEBHOOK_API_TOKEN
    }
    # print(data)
    response= requests.post(url, json=data)
    # pprint(response)
    return response.json()


def send_first_message_to_manager(telegram_id, text, first_name, last_name, username):
    """Отправляет первое сообщение менеджеру в чат на сайте для регистрации пользователя"""
    url=f'http://{IP_CHAT_ROOM}/api/webhook/client-message'
    data={
        'telegram_id': telegram_id,
        'text': text,
        'token': WEBHOOK_API_TOKEN,
        'first_name': first_name,
        'last_name': last_name,
        'username': username
    }
    response= requests.post(url, json=data)
    return response.json()

def send_message_to_manager(telegram_id, text):
    """Отправляет сообщение менеджеру в чат на сайте"""
    url=f'http://{IP_CHAT_ROOM}/api/webhook/client-message'
    data={
        'telegram_id': telegram_id,
        'text': text,
        'token': WEBHOOK_API_TOKEN
    }
    response= requests.post(url, json=data)
    return response.json()


def send_file_message_to_manager(telegram_id, name_file, data_file, text='Файл отправлен'):
    """Отправляет файл менеджеру в чат на сайте"""
    url=f'http://{IP_CHAT_ROOM}/api/webhook/client-message'
    data={
        'telegram_id': telegram_id,
        'text': text,
        'token': WEBHOOK_API_TOKEN,
        'file': {
            'name': name_file,
            'data': data_file,
        },
    }
    response= requests.post(url, json=data)
    return response.json()

if __name__ == '__main__':
    # path='/8 марта 204д - 116 (16 этаж)/Где_роутер_и_щиток.heic'
    # a=convert_heic_to_jpg(path)
    # print(a)
    # send_message_to_client(400923372, 'Здравствуйте! Это сообщение от клиента через API.')
    send_message_to_manager(400923372, 'Здравствуйте! Это сообщение от менеджера через API. от бота')
    # # a=get_path_to_file_instruction(url, returnIsUrl=True)
    # # pprint(a)
    # # a=get_path_to_file_info(url, returnIsUrl=True)
    # a=get_path_to_file_info(url, returnIsUrl=False)
    # pprint(a)
    # MAIN_PATH='✅ИНСТРУКЦИИ и ИНФО по крвартирам'
    # b=['/ИНФО ПО КВАРТИРАМ/Азина 22 корп.2 -198 (11 этаж)/Как выглядит квартира.mov',
    #  '/ИНФО ПО КВАРТИРАМ/Азина 22 корп.2 -198 (11 этаж)/Номер охраны.JPG',
    #  '/ИНФО ПО КВАРТИРАМ/Азина 22 корп.2 -198 (11 этаж)/Про ТВ в спальне.JPG']

